from fastapi import FastAPI, UploadFile, File, Form, Request, BackgroundTasks, Query
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse
from typing import List, Optional, Dict
import fitz
from pydantic import BaseModel
import re
from constants import ada_002_model, gpt_4o_model, GPT_4o_MODEL_NAME
import polars as pl
from sklearn.metrics.pairwise import cosine_similarity
from openpyxl import load_workbook
from pathlib import Path
from base64 import b64encode
from docx import Document as DocxDocument
from email import policy
from email.parser import BytesParser
import os
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.middleware.cors import CORSMiddleware
from fastapi.exceptions import HTTPException
from traceback import print_exc
from datetime import datetime, timezone, timedelta
from asyncio import sleep
from uuid import uuid4
from io import BytesIO
from tenacity import (
    retry,
    stop_after_attempt,
    wait_random_exponential,
)
from middlewares.logging_middleware import log_requests
import logging
from pyzbar.pyzbar import decode
from PIL import Image
import cv2
import numpy as np
import difflib
import webbrowser
import pytesseract

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s | %(levelname)s | %(filename)s:%(lineno)d | %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S',
    handlers=[
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

global_sessions = {}
df = pl.DataFrame()
logger.info("Starting the application...")
logger.info(df)

app = FastAPI()
app.middleware("http")(log_requests)
app.mount("/static", StaticFiles(directory="static"), name='static')
templates = Jinja2Templates("templates")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

uploads_dir = Path("static/uploads")
uploads_dir.mkdir(exist_ok=True, parents=True)

# Models
class FileUploadResponse(BaseModel):
    id: str
    filename: str
    content_type: str
    size: int

class URLAddRequest(BaseModel):
    url: str

class URLResponse(BaseModel):
    id: str
    url: str

class ChatMessage(BaseModel):
    user: str
    message: str

class QRCodeResponse(BaseModel):
    qr_codes: List[Dict]
    urls_opened: List[str]

# Core Functions
def recursive_character_splitter(text: str, chunk_size: int = 500, overlapping_size: int = 100, separators: Optional[List[str]] = None) -> List[str]:
    if separators is None:
        separators = ["\n\n", "\n", " ", ""]

    def _split_text(text: str, separators: List[str]) -> List[str]:
        final_chunks = []
        separator = separators[-1]
        new_separators = []
        for i, _s in enumerate(separators):
            _separator = re.escape(_s)
            if _s == "":
                separator = _s
                break
            if re.search(_separator, text):
                separator = _s
                new_separators = separators[i + 1:]
                break

        _separator = re.escape(separator)
        splits = re.split(_separator, text)

        _good_splits = []
        for s in splits:
            if len(s) < chunk_size:
                _good_splits.append(s)
            else:
                if _good_splits:
                    final_chunks.append(separator.join(_good_splits))
                    _good_splits = []
                if not new_separators:
                    final_chunks.append(s)
                else:
                    other_info = _split_text(s, new_separators)
                    final_chunks.extend(other_info)
        if _good_splits:
            final_chunks.append(separator.join(_good_splits))
        return final_chunks

    chunks = _split_text(text.lower(), separators)
    overlapping_substrings = []
    for i in range(len(chunks) - 1):
        overlapping_substrings.append(chunks[i].strip())
        overlap = chunks[i][-overlapping_size:] + chunks[i+1][:chunk_size - overlapping_size]
        overlapping_substrings.append(overlap.strip())
    overlapping_substrings.append(chunks[-1].strip())
    return overlapping_substrings

def create_embeddings(chunks: List[str], text: str, file_name: str) -> List[dict]:
    documents = []
    chunks = [c.strip() for c in chunks if c.strip()]
    embeddings = ada_002_model.embeddings.create(input=chunks, model="text-data-002").data
    for i, e in enumerate(embeddings):
        documents.append({
            "chunk": chunks[i].strip(),
            "embedding": e.embedding,
            "text": text,
            "file_name": file_name,
        })
    return documents

# Document Processing Functions
def process_pdf(file_path: Path) -> str:
    pdf_document = fitz.open(file_path)
    text = ""
    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        t = page.get_text("text")
        if t:
            text += t + "\n\n"
        else:
            img_bytes = page.get_pixmap(dpi=400).tobytes()
            t = process_image(img_bytes)
            if t:
                text += t + "\n\n"
    
    chunks = recursive_character_splitter(text, chunk_size=800, overlapping_size=200)
    pdf_documents = create_embeddings(chunks, text, file_path.name)
    return pdf_documents

def process_image(file_path: Path | bytes) -> str:
    if isinstance(file_path, Path):
        file_data = file_path.read_bytes()
        img_b64 = b64encode(file_data).decode("utf-8")
    elif isinstance(file_path, bytes):
        img_b64 = b64encode(file_path).decode("utf-8")

    messages = [
        {
            "role": "system",
            "content": [{"type": "text", "text": "You are an excellent image transcription AI assistant. You will help the user by transcribing the text from the image.Put the text in the <transcription> tags. If there is no text just say so within transcription tags."}],
        },
        {
            "role": "user",
            "content": [
                {"type": "text", "text": "please transcribe the text from this image"},
                {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_b64}"}}
            ],
        },
    ]

    response = gpt_4o_model.chat.completions.create(
        model=GPT_4o_MODEL_NAME,
        messages=messages,
        temperature=0.0,
        max_tokens=4078,
    )

    llm_answer = response.choices[0].message.content
    transcriptions = re.findall(r"<transcription>(.*?)</transcription>", llm_answer, re.DOTALL)
    if transcriptions:
        text = transcriptions[0]
        if isinstance(file_path, bytes):
            return text
        chunks = recursive_character_splitter(text, chunk_size=800, overlapping_size=200)
        file_name = file_path.name
        image_documents = create_embeddings(chunks, text, file_name)
        return image_documents
    else:
        return "Image transcription failed."

def process_xlsx(file_path: Path) -> str:
    wb = load_workbook(file_path)
    text = ""
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        for row in ws.iter_rows(values_only=True):
            for cell in row:
                if cell:
                    text += str(cell) + " "
            text += "\n\n"

    chunks = recursive_character_splitter(text, chunk_size=800, overlapping_size=200)
    xlsx_documents = create_embeddings(chunks, text, file_path)
    return xlsx_documents

def process_docx(file_path: Path) -> str:
    doc = DocxDocument(file_path)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    
    text = "\n".join(text)
    chunks = recursive_character_splitter(text, chunk_size=800, overlapping_size=200, separators=["\n\n"])
    docx_documents = create_embeddings(chunks, text, file_path.name)
    return docx_documents

def process_eml(file_path: Path) -> str:
    with file_path.open("rb") as f:
        msg = BytesParser(policy=policy.default).parse(f)
        text = msg.get_body(preferencelist=("plain", "html")).get_content()

    chunks = recursive_character_splitter(text, chunk_size=800, overlapping_size=200)
    eml_documents = create_embeddings(chunks, text, file_path)
    return eml_documents

# QR Code Functions
def extract_qr_codes_from_image(image_path: Path | bytes) -> List[str]:
    """Extract QR codes from an image file or bytes"""
    if isinstance(image_path, Path):
        img = Image.open(image_path)
    elif isinstance(image_path, bytes):
        img = Image.open(BytesIO(image_path))
    
    open_cv_image = np.array(img)
    if len(open_cv_image.shape) == 3:  # Convert to grayscale if color image
        open_cv_image = cv2.cvtColor(open_cv_image, cv2.COLOR_RGB2GRAY)
    
    _, thresh = cv2.threshold(open_cv_image, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    decoded_objects = decode(thresh)
    
    return [obj.data.decode('utf-8') for obj in decoded_objects]

def extract_qr_codes_from_pdf(pdf_path: Path) -> List[dict]:
    """Extract QR codes from a PDF file"""
    pdf_document = fitz.open(pdf_path)
    results = []
    
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        pix = page.get_pixmap(dpi=300)
        img_data = np.frombuffer(pix.samples, dtype=np.uint8).reshape((pix.height, pix.width, pix.n))
        
        gray = cv2.cvtColor(img_data, cv2.COLOR_RGB2GRAY)
        _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

        decoded_objects = decode(thresh)
        if decoded_objects:
            for obj in decoded_objects:
                results.append({
                    "page": page_num + 1,
                    "data": obj.data.decode('utf-8'),
                    "type": obj.type
                })
    
    pdf_document.close()
    return results

# Document Comparison Functions
def compare_documents(text1: str, text2: str) -> str:
    """Compare two text documents and return HTML diff"""
    differ = difflib.HtmlDiff()
    html_diff = differ.make_file(
        text1.splitlines(), 
        text2.splitlines(),
        fromdesc="Document 1",
        todesc="Document 2"
    )
    
    css = """
    <style>
        table.diff {font-family:Courier; border:medium; width:100%;}
        .diff_header {background-color:#e0e0e0}
        td.diff_header {text-align:right}
        .diff_next {background-color:#c0c0c0}
        .diff_add {background-color:#aaffaa}
        .diff_chg {background-color:#ffff77}
        .diff_sub {background-color:#ffaaaa}
    </style>
    """
    
    return css + html_diff

# Text Extraction Helpers
def extract_text_from_pdf(pdf_path: Path) -> str:
    try:
        text = ""
        pdf_document = fitz.open(pdf_path)
        for page in pdf_document:
            text += page.get_text()
        return text
    except Exception as e:
        raise ValueError(f"PDF error: {str(e)}")

def extract_text_from_image(image_path: Path) -> str:
    try:
        img = Image.open(image_path)
        return pytesseract.image_to_string(img)
    except Exception as e:
        raise ValueError(f"Image error: {str(e)}")

def extract_text_from_excel(excel_path: Path) -> str:
    try:
        wb = load_workbook(excel_path, data_only=True)
        return '\n'.join(
            ' '.join(str(cell) for cell in row)
            for sheet in wb
            for row in sheet.iter_rows(values_only=True)
        )
    except Exception as e:
        raise ValueError(f"Excel error: {str(e)}")

def extract_text_from_word(word_path: Path) -> str:
    try:
        doc = DocxDocument(word_path)
        return '\n'.join(para.text for para in doc.paragraphs)
    except Exception as e:
        raise ValueError(f"Word error: {str(e)}")

def extract_text_from_eml(eml_path: Path) -> str:
    try:
        with open(eml_path, 'rb') as fp:
            msg = BytesParser(policy=policy.default).parse(fp)
        text_parts = []
        for part in msg.walk():
            if part.get_content_type() == 'text/plain':
                text_parts.append(part.get_payload(decode=True).decode('utf-8', errors='replace'))
        return '\n'.join(text_parts)
    except Exception as e:
        raise ValueError(f"EML error: {str(e)}")

def extract_text_from_file(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == '.pdf':
        return extract_text_from_pdf(file_path)
    elif suffix in ('.jpg', '.jpeg', '.png'):
        return extract_text_from_image(file_path)
    elif suffix in ('.xlsx', '.xls'):
        return extract_text_from_excel(file_path)
    elif suffix == '.docx':
        return extract_text_from_word(file_path)
    elif suffix == '.eml':
        return extract_text_from_eml(file_path)
    elif suffix == '.txt':
        return file_path.read_text()
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

# API Endpoints
@app.post("/extract-qr-codes", response_model=QRCodeResponse)
async def extract_qr_codes(file: UploadFile = File(...)):
    """Endpoint to extract QR codes from a file"""
    try:
        file_location = uploads_dir / file.filename
        file_location.write_bytes(file.file.read())
        
        if file.filename.lower().endswith('.pdf'):
            results = extract_qr_codes_from_pdf(file_location)
        elif file.filename.lower().endswith(('.jpg', '.jpeg', '.png')):
            results = extract_qr_codes_from_image(file_location)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type for QR code extraction")
        
        file_location.unlink()
        
        urls_opened = []
        for result in results:
            if isinstance(result, dict):
                data = result['data']
            else:
                data = result
                
            if data.startswith(('http://', 'https://')):
                webbrowser.open(data)
                urls_opened.append(data)
        
        return {
            "qr_codes": results,
            "urls_opened": urls_opened
        }
        
    except Exception as e:
        logger.error(f"Error extracting QR codes: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error extracting QR codes: {str(e)}")

@app.post("/compare-documents")
async def compare_documents_endpoint(
    file1: UploadFile = File(...),
    file2: UploadFile = File(...)
):
    """Endpoint to compare two documents"""
    try:
        # Save and process first file
        file1_location = uploads_dir / file1.filename
        file1_location.write_bytes(file1.file.read())
        text1 = extract_text_from_file(file1_location)
        file1_location.unlink()
        
        # Save and process second file
        file2_location = uploads_dir / file2.filename
        file2_location.write_bytes(file2.file.read())
        text2 = extract_text_from_file(file2_location)
        file2_location.unlink()
        
        html_diff = compare_documents(text1, text2)
        return HTMLResponse(content=html_diff)
        
    except Exception as e:
        logger.error(f"Error comparing documents: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error comparing documents: {str(e)}")

# Existing endpoints (updated with proper typing and error handling)
@app.post("/session")
async def get_session(background_tasks: BackgroundTasks, request: Request):
    data = await request.json()
    session = data.get("session")
    global_sessions[session] = {}
    removal_time = datetime.now(tz=timezone.utc) + timedelta(seconds=60*30)
    background_tasks.add_task(remove_session, session, removal_time)
    return {"session": session}

async def remove_session(session: str, removal_time: datetime):
    while datetime.now(tz=timezone.utc) < removal_time:
        await sleep(60 * 5)
    
    if session in global_sessions:
        del global_sessions[session]
        logger.info(f"Session {session} removed at {datetime.now(tz=timezone.utc)}")

@app.post("/upload/files", response_model=List[FileUploadResponse])
async def upload_files(request: Request, files: List[UploadFile] = File(...)):
    try:
        data = await request.form()
        session = data.get("session")

        session_uploaded_files = global_sessions.get(session, {}).get("uploaded_files", [])
        for file in files:
            file_id = str(uuid4())
            file_details = {
                "id": file_id,
                "filename": file.filename,
                "content_type": file.content_type,
                "size": os.fstat(file.file.fileno()).st_size,
            }
            session_uploaded_files.append(file_details)
        
        result = process_documents(files, session)
        global_sessions[session] = {"uploaded_files": session_uploaded_files}
        logger.info(df)
        return session_uploaded_files
    except Exception as e:
        logger.info(e)
        logger.info(print_exc())
        raise HTTPException(status_code=500, detail="Failed to process files")

def process_documents(files: List[UploadFile], session: str) -> str:
    global df
    mapping = {
        ".pdf": process_pdf,
        ".jpg": process_image,
        ".jpeg": process_image,
        ".png": process_image,
        ".xlsx": process_xlsx,
        ".xls": process_xlsx,
        ".docx": process_docx,
        ".eml": process_eml,
    }

    documents = []
    for file in files:
        file_location = uploads_dir / file.filename
        file_location.write_bytes(file.file.read())
        filename = Path(file.filename)
        if filename.suffix not in mapping:
            return "Unsupported file type. Please upload a PDF, Image, Excel, Word, or EML file."
        
        document = mapping[filename.suffix](file_location)
        documents.extend(document)
        file_location.unlink()
    
    try:
        tmp_df = pl.DataFrame(documents)
        tmp_df = tmp_df.with_columns(pl.lit(session).alias("session")).with_columns(pl.col("chunk").str.split(by=" ").list.len().alias("len_of_chunks"))
        df = pl.concat([df, tmp_df])
        df = df.sort("chunk").set_sorted("chunk")
        return "Documents successfully processed and embedded."
    except Exception as e:
        print(e)
        return "Documents processing failed."

@app.post("/upload/xlsx", response_model=List[FileUploadResponse])
async def upload_xlsx(request: Request, background_task: BackgroundTasks, files: List[UploadFile] = File(...)):
    data = await request.form()
    session = data.get("session")

    session_checkpoint_files = global_sessions.get(session, {}).get("checkpoint_files", [])
    results = []

    for file in files:
        if file.filename.endswith(".xlsx"):
            file_id = str(uuid4())
            file_details = {
                "id": file_id,
                "filename": file.filename,
                "content_type": file.content_type,
                "size": os.fstat(file.file.fileno()).st_size,
            }
            session_checkpoint_files.append(file_details)

            contents = await file.read()
            df = pl.read_excel(BytesIO(contents))
            columns = df.columns
            messages = df[columns[0]].to_list()

            for i, message in enumerate(messages):
                answer = ask_question(message, session, checklist=True)
                results.append({"message": message, "answer": answer})
    
    result_df = pl.DataFrame(results)
    filename = uploads_dir / f"results_{datetime.now(tz=timezone.utc).strftime('%Y%m%d%H%M%S')}.csv"
    result_df.write_csv(filename)
    background_task.add_task(remove_file_after_download, filename)
    return FileResponse(filename, media_type="text/csv", filename=filename.name)

def remove_file_after_download(file_path: Path):
    file_path.unlink()

@app.post("/add/url", response_model=List[URLResponse])
async def add_url(url_data: URLAddRequest, request: Request):
    data = await request.json()
    session = data.get("session")

    if not url_data.url.startswith("http"):
        raise HTTPException(status_code=400, detail="Invalid URL format")
    if len(url_data.url) > 1000:
        raise HTTPException(status_code=400, detail="URL too long")
    if len(url_data.url.strip()) == 0:
        raise HTTPException(status_code=400, detail="URL cannot be empty")
    
    session_uploaded_urls = global_sessions.get(session, {}).get("uploaded_urls", [])
    url_id = str(uuid4())
    url_entry = {"id": url_id, "url": url_data.url}
    session_uploaded_urls.append(url_entry)
    global_sessions[session] = {"uploaded_urls": session_uploaded_urls}
    return session_uploaded_urls

@app.post("/files", response_model=List[FileUploadResponse])
async def list_files(request: Request):
    data = await request.json()
    session = data.get("session")
    uploaded_files = global_sessions.get(session, {}).get("uploaded_files", [])
    return uploaded_files

@app.post("/urls", response_model=List[URLResponse])
async def list_urls(request: Request):
    data = await request.json()
    session = data.get("session")
    uploaded_urls = global_sessions.get(session, {}).get("uploaded_urls", [])
    return uploaded_urls

@app.delete("/files/{file_id}")
async def delete_file(file_id: str, request: Request):
    data = await request.json()
    session = data.get("session")

    global df
    uploaded_files = global_sessions.get(session, {}).get("uploaded_files", [])
    uploaded_files_after_delete = []

    for file in uploaded_files:
        if file["id"] != file_id:
            uploaded_files_after_delete.append(file)
        else:
            df = df.filter((pl.col("file_name") != file["filename"]) & (pl.col("session") == session))

    global_sessions[session] = {"uploaded_files": uploaded_files_after_delete}
    return {"message": "File deleted successfully"}

@app.delete("/urls/{url_id}")
async def delete_url(url_id: str, request: Request):
    data = await request.json()
    session = data.get("session")
    uploaded_urls = global_sessions.get(session, {}).get("uploaded_urls", [])
    uploaded_urls = [url for url in uploaded_urls if url["id"] != url_id]
    global_sessions[session] = {"uploaded_urls": uploaded_urls}
    return {"message": "URL deleted successfully"}

@app.post("/chat/send", response_model=ChatMessage)
async def send_message(message: ChatMessage, request: Request):
    data = await request.json()
    session = data.get("session")
    answer = ask_question(message.message, session)
    return {"user": "bot", "message": answer, "session": session}

def ask_question(question: str, session, checklist=False):
    global df
    if df.height == 0:
        return "Please process the documents first."
    
    df1 = df.filter(pl.col("session") == session).sort("chunk").set_sorted("chunk")
    if df1.height == 0:
        return "Please process the documents first."
    
    train_embeddings = df1.filter(pl.col("session") == session)["embedding"].to_list()
    q_ems = ada_002_model.embeddings.create(input=[question.lower()], model="text-data-002").data
    q_em = [q_ems[0].embedding]
    semantic_search_results = cosine_similarity(train_embeddings, q_em).flatten()
    top_5_indices = semantic_search_results.argsort()[-5:][::-1]

    context = []
    for i in top_5_indices:
        row = df1.row(i, named=True)
        document_name = row["file_name"]
        chunk = row["chunk"]
        context.append({"document_name": document_name, "chunk": chunk})
    
    system_prompt = """You are an AI assistant, you will answer user questions based on the provided document only. If you can't find the answer in the documents, you will say so.

Here's how you will work:
1. You will go through all the documents one by one to find the answer to the user's question.
2. You will give a user-facing answer in <answer> tags.
3. If you can't find the answer in the documents, you will indicate that you don't know in <answer> tags.
4. No Hallucination and Be Honest."""

    if checklist:
        system_prompt += "5. User wants to do entity extraction like getting the names of the people, places, etc. from the documents. You will provide the answer in <answer> tags."

    system_prompt += "\n\nYou will use markdown format for the answer."
    user_prompt = "<documents>"
    for i, c in enumerate(context):
        user_prompt += f"""
    <document index={i}>
        <document_content>
            {c["chunk"]}
        </document_content>
    </document>
"""
    user_prompt += "</documents>\n\n"
    user_prompt += f"<question>{question}</question>"

    messages = [
        {
            "role": "system",
            "content": [{"type": "text", "text": system_prompt}],
        },
        {
            "role": "user",
            "content": [{"type": "text", "text": user_prompt}],
        },
    ]

    @retry(wait=wait_random_exponential(min=5, max=150), stop=stop_after_attempt(10))
    def completion_with_backoff():
        return gpt_4o_model.chat.completions.create(
            model=GPT_4o_MODEL_NAME,
            messages=messages,
            temperature=0.1,
            max_tokens=4078,
        )

    response = completion_with_backoff()
    llm_answer = response.choices[0].message.content

    final_answer = re.findall(r"<answer>(.*?)</answer>", llm_answer, re.DOTALL)
    if final_answer:
        return final_answer[0].strip()
    else:
        return llm_answer.strip()

@app.get("/ping")
async def root():
    return {"message": "Welcome to the DocChat API!"}

@app.get("/", response_class=HTMLResponse)
async def read_root(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True)